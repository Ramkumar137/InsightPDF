"""
Export service for generating downloadable files in multiple formats
"""
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from datetime import datetime
from app.models.summary import Summary

class ExportService:
    """Service for exporting summaries to different file formats"""
    
    @staticmethod
    def export_to_txt(summary: Summary) -> bytes:
        """
        Export summary to plain text format.
        
        Args:
            summary: Summary model instance
            
        Returns:
            Bytes content of text file
        """
        content = []
        content.append(f"PDF SUMMARY REPORT")
        content.append(f"=" * 50)
        content.append(f"\nFile(s): {', '.join(summary.file_names)}")
        content.append(f"Context: {summary.context_type.title()}")
        content.append(f"Generated: {summary.created_at.strftime('%Y-%m-%d %H:%M:%S')}")
        content.append(f"\n{'=' * 50}\n")
        
        content.append(f"\n📋 OVERVIEW\n{'-' * 50}")
        content.append(summary.overview)
        
        content.append(f"\n\n💡 KEY INSIGHTS\n{'-' * 50}")
        content.append(summary.key_insights)
        
        if summary.risks:
            content.append(f"\n\n⚠️  RISKS & CHALLENGES\n{'-' * 50}")
            content.append(summary.risks)
        
        if summary.recommendations:
            content.append(f"\n\n✅ RECOMMENDATIONS\n{'-' * 50}")
            content.append(summary.recommendations)
        
        content.append(f"\n\n{'=' * 50}")
        content.append(f"End of Report")
        
        text_content = "\n".join(content)
        return text_content.encode('utf-8')
    
    @staticmethod
    def export_to_pdf(summary: Summary) -> bytes:
        """
        Export summary to PDF format using ReportLab.
        
        Args:
            summary: Summary model instance
            
        Returns:
            Bytes content of PDF file
        """
        buffer = io.BytesIO()
        
        # Create PDF document
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )
        
        # Container for PDF elements
        elements = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor='#2563eb',
            spaceAfter=30,
            alignment=1  # Center
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor='#1e40af',
            spaceAfter=12,
            spaceBefore=20
        )
        
        # Title
        title = Paragraph("PDF Summary Report", title_style)
        elements.append(title)
        elements.append(Spacer(1, 0.2 * inch))
        
        # Metadata
        metadata_style = styles['Normal']
        meta_text = f"""
        <b>File(s):</b> {', '.join(summary.file_names)}<br/>
        <b>Context:</b> {summary.context_type.title()}<br/>
        <b>Generated:</b> {summary.created_at.strftime('%Y-%m-%d %H:%M:%S')}
        """
        elements.append(Paragraph(meta_text, metadata_style))
        elements.append(Spacer(1, 0.3 * inch))
        
        # Overview section
        elements.append(Paragraph("📋 Overview", heading_style))
        elements.append(Paragraph(summary.overview, styles['Normal']))
        elements.append(Spacer(1, 0.2 * inch))
        
        # Key Insights
        elements.append(Paragraph("💡 Key Insights", heading_style))
        elements.append(Paragraph(summary.key_insights, styles['Normal']))
        elements.append(Spacer(1, 0.2 * inch))
        
        # Risks (if available)
        if summary.risks:
            elements.append(Paragraph("⚠️ Risks & Challenges", heading_style))
            elements.append(Paragraph(summary.risks, styles['Normal']))
            elements.append(Spacer(1, 0.2 * inch))
        
        # Recommendations (if available)
        if summary.recommendations:
            elements.append(Paragraph("✅ Recommendations", heading_style))
            elements.append(Paragraph(summary.recommendations, styles['Normal']))
        
        # Build PDF
        doc.build(elements)
        
        # Get PDF bytes
        pdf_bytes = buffer.getvalue()
        buffer.close()
        
        return pdf_bytes
    
    @staticmethod
    def export_to_docx(summary: Summary) -> bytes:
        """
        Export summary to DOCX format using python-docx.
        
        Args:
            summary: Summary model instance
            
        Returns:
            Bytes content of DOCX file
        """
        doc = Document()
        
        # Title
        title = doc.add_heading('PDF Summary Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        doc.add_paragraph()
        meta = doc.add_paragraph()
        meta.add_run('File(s): ').bold = True
        meta.add_run(', '.join(summary.file_names))
        
        meta = doc.add_paragraph()
        meta.add_run('Context: ').bold = True
        meta.add_run(summary.context_type.title())
        
        meta = doc.add_paragraph()
        meta.add_run('Generated: ').bold = True
        meta.add_run(summary.created_at.strftime('%Y-%m-%d %H:%M:%S'))
        
        doc.add_paragraph()
        doc.add_paragraph('_' * 50)
        
        # Overview
        doc.add_heading('📋 Overview', 1)
        doc.add_paragraph(summary.overview)
        
        # Key Insights
        doc.add_heading('💡 Key Insights', 1)
        doc.add_paragraph(summary.key_insights)
        
        # Risks (if available)
        if summary.risks:
            doc.add_heading('⚠️ Risks & Challenges', 1)
            doc.add_paragraph(summary.risks)
        
        # Recommendations (if available)
        if summary.recommendations:
            doc.add_heading('✅ Recommendations', 1)
            doc.add_paragraph(summary.recommendations)
        
        # Footer
        doc.add_paragraph()
        doc.add_paragraph('_' * 50)
        footer = doc.add_paragraph('End of Report')
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save to bytes buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()
        buffer.close()
        
        return docx_bytes

# Global export service instance
export_service = ExportService()